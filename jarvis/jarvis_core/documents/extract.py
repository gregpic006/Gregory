"""Extraction de texte a partir d'un fichier.

Chaque extracteur renvoie des **segments localises**: un morceau de texte et
l'endroit d'ou il vient (page 4, section « Resiliation »).  Sans ce reperage,
JARVIS pourrait citer un document sans pouvoir dire ou regarder — ce qui
revient a demander qu'on le croie sur parole.
"""

from __future__ import annotations

import contextlib
import logging
from dataclasses import dataclass
from pathlib import Path

from jarvis_core.errors import DocumentError

logger = logging.getLogger(__name__)

#: Extensions reconnues.  Tout le reste est refuse explicitement plutot que
#: lu comme du texte au hasard (un .xlsx lu en UTF-8 produit du charabia).
SUPPORTED_SUFFIXES = frozenset({".txt", ".md", ".markdown", ".pdf", ".docx"})

#: Au-dela, on refuse: un fichier de 200 Mo ferait exploser la memoire et
#: n'apporterait rien d'exploitable a une recherche.
MAX_BYTES = 40 * 1024 * 1024


@dataclass(frozen=True)
class Segment:
    """Un morceau de document, avec l'endroit d'ou il vient."""

    text: str
    locator: str = ""
    """Repere lisible par un humain: « page 3 », « section Resiliation »."""


def is_supported(path: Path | str) -> bool:
    return Path(path).suffix.lower() in SUPPORTED_SUFFIXES


def extract(path: Path | str) -> list[Segment]:
    """Extrait les segments d'un fichier.

    Raises:
        DocumentError: fichier absent, trop gros, format non gere ou illisible.
    """
    file = Path(path)
    if not file.exists():
        raise DocumentError(
            f"fichier absent: {file}",
            user_message=f"Je ne trouve pas le fichier {file.name}.",
        )
    suffix = file.suffix.lower()
    if suffix not in SUPPORTED_SUFFIXES:
        formats = ", ".join(sorted(SUPPORTED_SUFFIXES))
        raise DocumentError(
            f"format non gere: {suffix}",
            user_message=(
                f"Je ne sais pas lire un fichier {suffix or 'sans extension'}. "
                f"Formats geres: {formats}."
            ),
        )
    size = file.stat().st_size
    if size > MAX_BYTES:
        raise DocumentError(
            f"fichier trop volumineux: {size} octets",
            user_message=(
                f"{file.name} fait {size // (1024 * 1024)} Mo, "
                f"c'est au-dela de la limite de {MAX_BYTES // (1024 * 1024)} Mo."
            ),
        )

    if suffix == ".pdf":
        segments = _extract_pdf(file)
    elif suffix == ".docx":
        segments = _extract_docx(file)
    else:
        segments = _extract_text(file)

    kept = [s for s in segments if s.text.strip()]
    if not kept:
        raise DocumentError(
            f"aucun texte extrait de {file}",
            user_message=(
                f"{file.name} ne contient aucun texte lisible. "
                "S'il s'agit d'un PDF scanne, il faudrait le passer par une "
                "reconnaissance de caracteres d'abord."
            ),
        )
    return kept


def _extract_text(file: Path) -> list[Segment]:
    """Fichier texte ou Markdown, decoupe par titres quand il y en a."""
    try:
        raw = file.read_text(encoding="utf-8")
    except UnicodeDecodeError:
        # Beaucoup de fichiers venant de Windows sont en cp1252.
        try:
            raw = file.read_text(encoding="cp1252")
        except (UnicodeDecodeError, OSError) as exc:
            raise DocumentError(
                f"encodage illisible: {file}",
                user_message=f"Je n'arrive pas a decoder {file.name}.",
            ) from exc
    except OSError as exc:
        raise DocumentError(
            f"lecture impossible: {file}",
            user_message=f"Je n'arrive pas a ouvrir {file.name}.",
        ) from exc

    if file.suffix.lower() not in {".md", ".markdown"}:
        return [Segment(text=raw)]

    segments: list[Segment] = []
    heading = ""
    buffer: list[str] = []
    for line in raw.splitlines():
        if line.startswith("#"):
            if buffer:
                segments.append(Segment(text="\n".join(buffer), locator=heading))
                buffer = []
            heading = line.lstrip("#").strip()
            # On garde le titre dans le texte — il porte du sens pour la
            # recherche — mais sans les dieses, qui ne sont que du balisage
            # et se retrouveraient tels quels dans les extraits affiches.
            buffer.append(heading)
            continue
        buffer.append(line)
    if buffer:
        segments.append(Segment(text="\n".join(buffer), locator=heading))
    return segments


def _extract_pdf(file: Path) -> list[Segment]:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependance declaree
        raise DocumentError(
            "pypdf absent",
            user_message="Le lecteur de PDF n'est pas installe. Relance setup.ps1.",
        ) from exc

    try:
        reader = PdfReader(str(file))
        if reader.is_encrypted:
            # Un PDF protege par mot de passe vide s'ouvre parfois ainsi.
            with contextlib.suppress(Exception):
                reader.decrypt("")
        pages = list(reader.pages)
    except DocumentError:
        raise
    except Exception as exc:  # noqa: BLE001 - pypdf leve des exceptions variees
        raise DocumentError(
            f"pdf illisible: {file} ({exc})",
            user_message=(
                f"Je n'arrive pas a lire {file.name}. "
                "Le fichier est peut-etre protege par un mot de passe ou abime."
            ),
        ) from exc

    segments: list[Segment] = []
    for number, page in enumerate(pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # noqa: BLE001 - une page abimee n'annule pas le reste
            logger.warning("page %s de %s illisible: %s", number, file.name, exc)
            continue
        segments.append(Segment(text=text, locator=f"page {number}"))
    return segments


def _extract_docx(file: Path) -> list[Segment]:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependance declaree
        raise DocumentError(
            "python-docx absent",
            user_message="Le lecteur de documents Word n'est pas installe. Relance setup.ps1.",
        ) from exc

    try:
        document = docx.Document(str(file))
    except Exception as exc:  # noqa: BLE001 - python-docx leve des exceptions variees
        raise DocumentError(
            f"docx illisible: {file} ({exc})",
            user_message=f"Je n'arrive pas a lire {file.name}. Le fichier est peut-etre abime.",
        ) from exc

    segments: list[Segment] = []
    heading = ""
    buffer: list[str] = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if paragraph.style is not None and str(paragraph.style.name).startswith("Heading"):
            if buffer:
                segments.append(Segment(text="\n".join(buffer), locator=heading))
                buffer = []
            heading = text
        buffer.append(text)
    if buffer:
        segments.append(Segment(text="\n".join(buffer), locator=heading))

    # Les tableaux portent souvent l'essentiel (montants, echeances).
    for index, table in enumerate(document.tables, start=1):
        rows = [
            " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            for row in table.rows
        ]
        rows = [r for r in rows if r]
        if rows:
            segments.append(Segment(text="\n".join(rows), locator=f"tableau {index}"))
    return segments
